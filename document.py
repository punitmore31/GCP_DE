import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import yaml


def set_font_style(run, size_pt, is_bold=False, color_rgb=RGBColor(0, 0, 0)):
    font = run.font
    font.name = "Inter"
    font.size = Pt(size_pt)
    font.bold = is_bold
    font.color.rgb = color_rgb

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Inter")
    rFonts.set(qn("w:hAnsi"), "Inter")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def parse_complex_schedule(row, col_map):
    run_time_col = col_map["SCHEDULE_COLUMN"]
    if pd.isna(row[run_time_col]):
        return "N/A"
    return str(row[run_time_col])


def generate_process_overview(source_type):
    source_type_str = str(source_type).lower()
    if "sap" in source_type_str:
        return "This DAG extracts data from SAP source systems and loads it into the target tables."
    if "file" in source_type_str:
        return "This DAG extracts records from flat files and loads data into the target table."
    if "table" in source_type_str:
        return "This DAG processes data from source tables and loads results into target tables."
    return "This DAG processes and loads data into the target tables."


def generate_frequency_statement(source_type):
    source_type_str = str(source_type).lower()
    if "file" in source_type_str:
        return "Frequency - DAG checks for source files from GCS bucket"
    if "sap" in source_type_str:
        return (
            "Frequency - Triggered based on SAP data availability or a defined schedule"
        )
    return "Monthly Adhoc Run by Kaiku Users"


def determine_source_type(workflow_name, df_overview, col_map):
    df_overview.columns = df_overview.columns.str.strip()
    table_col = col_map["OVERVIEW_TABLE"]
    sap_col = col_map["OVERVIEW_SAP"]
    file_col = col_map["OVERVIEW_FILE"]
    if (
        table_col in df_overview.columns
        and workflow_name in df_overview[table_col].dropna().values
    ):
        return "Table"
    if (
        sap_col in df_overview.columns
        and workflow_name in df_overview[sap_col].dropna().values
    ):
        return "SAP"
    if (
        file_col in df_overview.columns
        and workflow_name in df_overview[file_col].dropna().values
    ):
        return "File"
    return "Unknown"


def get_all_workflow_details(wf_name, dataframes, config):
    details = {}
    col_map = config["COLUMN_MAPPINGS"]
    print(f"col_map : {col_map}")

    try:
        schedule_row = dataframes["schedule"].loc[wf_name]
        print(f"schedule_row : {schedule_row}")
        details["parsed_schedule"] = parse_complex_schedule(schedule_row, col_map)
        details["severity"] = schedule_row[col_map["SEVERITY"]]

        print(f"details dict : {details}")
    except (KeyError, IndexError):
        details["parsed_schedule"] = "Not found in schedule file"
        details["severity"] = "Low"

    try:
        details["task_config_path"] = dataframes["task_config"].loc[
            wf_name, col_map["TASK_CONFIG"]
        ]
    except (KeyError, IndexError):
        details["task_config_path"] = None

    source_type = determine_source_type(wf_name, dataframes["overview"], col_map)
    details["process_overview"] = generate_process_overview(source_type)
    details["frequency_statement"] = generate_frequency_statement(source_type)

    workflow_name_col = col_map["DAG_NAME_LOOKUP"]

    def find_data(df):
        data = df[df[workflow_name_col] == wf_name]
        if data.empty:
            short_name = wf_name.split(".")[-1]
            data = df[df[workflow_name_col].str.contains(short_name, na=False)]
        return data

    st_data = find_data(dataframes["source_target"])
    if not st_data.empty:
        details["sources"] = (
            st_data[col_map["SOURCE_TABLES"]].dropna().unique().tolist()
        )
        details["targets"] = (
            st_data[col_map["TARGET_TABLES"]].dropna().unique().tolist()
        )

    dc_data = find_data(dataframes["dag_config"])
    if not dc_data.empty:
        details["dag_config_path"] = (
            dc_data[col_map["DAG_CONFIG"]].dropna().unique().tolist() or [None]
        )[0]
        details["sql_dir_path"] = (
            dc_data[col_map["SQL_FILES"]].dropna().unique().tolist() or [None]
        )[0]

    return details


def generate_github_url(
    base_url, relative_path, repo_name_in_path, intended_extension=""
):
    if not relative_path:
        return "Path not found in CSV"
    clean_path = str(relative_path).replace("\\", "/")
    if clean_path.startswith(repo_name_in_path + "/"):
        clean_path = clean_path[len(repo_name_in_path) + 1 :]
    if intended_extension and not clean_path.endswith(intended_extension):
        clean_path += intended_extension
    return f"{base_url}/{clean_path}"


def create_runbook_doc(workflow_name, details, config):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    def setup_heading_style(level, size_pt, is_bold=False):
        style = doc.styles[f"Heading {level}"]
        font = style.font
        font.name = "Inter"
        font.size = Pt(size_pt)
        font.bold = is_bold
        font.color.rgb = RGBColor(0, 0, 0)
        font.italic = False

    setup_heading_style(1, 24, is_bold=True)
    for i in range(2, 8):
        setup_heading_style(i, 18)

    doc.add_heading(f"Runbook : {workflow_name}", level=1)
    p = doc.add_paragraph()
    set_font_style(p.add_run(f"Severity: {details.get('severity', 'Low')}\n"), 12)
    set_font_style(p.add_run(f"Business Impact: NA\n"), 12)
    set_font_style(
        p.add_run(f"Schedule: {details.get('parsed_schedule', 'N/A')}\n"), 12
    )
    set_font_style(p.add_run(f"GCP Project ID: {config['GCP_PROJECT_ID']}"), 12)
    doc.add_paragraph()
    doc.add_heading("Process Overview", level=2)
    p = doc.add_paragraph()
    set_font_style(p.add_run(details.get("process_overview", "")), 12)
    doc.add_paragraph()
    core_wf_name = workflow_name.split(".")[-1]
    screenshot_path = os.path.join(
        config["SCREENSHOTS_FOLDER_PATH"], f"{core_wf_name.lower()}.png"
    )
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(6.5))
    else:
        p = doc.add_paragraph()
        set_font_style(p.add_run(f"[Screenshot not found: {screenshot_path}]"), 11)
    doc.add_paragraph()
    doc.add_heading("Source and Target Table/File Details", level=3)
    p = doc.add_paragraph()
    set_font_style(p.add_run("Source Tables:"), 12, is_bold=True)
    for i, item in enumerate(details.get("sources", ["N/A"])):
        p = doc.add_paragraph()
        set_font_style(p.add_run(f"{i + 1}) {item}"), 12)
    p = doc.add_paragraph()
    set_font_style(p.add_run("Target Tables:"), 12, is_bold=True)
    for i, item in enumerate(details.get("targets", ["N/A"])):
        p = doc.add_paragraph()
        set_font_style(p.add_run(f"{i + 1}) {item}"), 12)
    doc.add_paragraph()
    doc.add_heading("DAG Config details", level=4)
    if details.get("dag_config_path"):
        url = generate_github_url(
            config["GITHUB_BLOB_URL"],
            details["dag_config_path"],
            config["REPO_NAME_IN_PATH"],
            ".yaml",
        )
        p = doc.add_paragraph()
        set_font_style(p.add_run("DAG Config: "), 12)
        add_hyperlink(p, url, url)
    if pd.notna(details.get("task_config_path")):
        url = generate_github_url(
            config["GITHUB_TREE_URL"],
            details["task_config_path"],
            config["REPO_NAME_IN_PATH"],
        )
        p = doc.add_paragraph()
        set_font_style(p.add_run("Task Config Path: "), 12)
        add_hyperlink(p, url, url)
    if details.get("sql_dir_path"):
        url = generate_github_url(
            config["GITHUB_TREE_URL"],
            details["sql_dir_path"],
            config["REPO_NAME_IN_PATH"],
        )
        p = doc.add_paragraph()
        set_font_style(p.add_run("SQL Files Path: "), 12)
        add_hyperlink(p, url, url)
    doc.add_paragraph()
    doc.add_heading("Execution Schedule", level=5)
    p = doc.add_paragraph()
    set_font_style(p.add_run(details.get("frequency_statement", "")), 12)
    doc.add_paragraph()
    doc.add_heading("Failure Scenario Handling", level=6)
    failure_texts = [
        "Rerun DAG/Failed task in case of failure.",
        "Check Airflow Logs for the specific failed task.",
        "Validate with audit table FDW_AUDIT.JOB_AUDIT using the workflow name to check the latest run status.",
    ]
    for text in failure_texts:
        p = doc.add_paragraph()
        set_font_style(p.add_run(text), 12)
    doc.add_paragraph()
    doc.add_heading("Escalation", level=7)
    p = doc.add_paragraph()
    set_font_style(p.add_run("On failure, drop an email to the associates below -"), 12)
    p = doc.add_paragraph()
    set_font_style(p.add_run(config["ESCALATION_EMAIL"]), 12)
    output_filename = f"Kaiku_Runbook_{workflow_name}.docx"
    output_path = os.path.join(config["OUTPUT_DIRECTORY"], output_filename)
    os.makedirs(config["OUTPUT_DIRECTORY"], exist_ok=True)
    doc.save(output_path)
    print(f"✅ Successfully generated runbook: {output_path}")


if __name__ == "__main__":
    CONFIG = {
        "SCHEDULE_DETAILS_PATH": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\input\Report_automation\Schedule_Schedule_details.csv",
        "SOURCE_TARGET_PATH": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\input\Report_automation\Runbook_details_Source_Target.csv",
        "DAG_CONFIG_PATH": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\input\Report_automation\RUN_BOOK_Dag_Config.csv",
        "TASK_CONFIG_PATH": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\input\Report_automation\RUN_BOOK_task_config.csv",
        "PROCESS_OVERVIEW_PATH": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\input\Report_automation\RUN_BOOK_File_Table_SAP.csv",
        "CODE_REPO_PATH": r"D:\Main\fdw-dags",
        "SCREENSHOTS_FOLDER_PATH": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\input\Report_automation\Screenshots",
        "OUTPUT_DIRECTORY": r"C:\Users\punitkumar.more\Documents\Elisa\gcp_de\AUTOMATION\output\DOC1",
        "GITHUB_BLOB_URL": "https://github.com/elisadatalake/fdw-dags/blob/main",
        "GITHUB_TREE_URL": "https://github.com/elisadatalake/fdw-dags/tree/main",
        "REPO_NAME_IN_PATH": "fdw-dags",
        "GCP_PROJECT_ID": "fdw-prod-5288",
        "ESCALATION_EMAIL": "elisa-migration-support@onixnet.com",
        "COLUMN_MAPPINGS": {
            "WORKFLOW_NAME": "Workflow_Name",
            "SCHEDULE_COLUMN": "Schedule",
            "SEVERITY": "Severity",
            "DAG_NAME_LOOKUP": "Workflow_Name",
            "SOURCE_TABLES": "Source_Tables",
            "TARGET_TABLES": "Target_Tables",
            "DAG_CONFIG": "DAG_CONFIG",
            "SQL_FILES": "SQL_FILES",
            "TASK_CONFIG": "Task_Config_Path",
            "OVERVIEW_TABLE": "table_source_workflow_name",
            "OVERVIEW_SAP": "sap_source_workflow_name",
            "OVERVIEW_FILE": "file_source_workflow_name",
        },
    }

    workflows_to_generate = [
        "wf_AR1_to_kaiku_stg",
        "wf_BASE_to_kaiku_stg",
        "wf_device_to_kaiku_stg",
        "wf_dynamo_kaiku_stg",
        "wf_FIRA_to_kaiku_stg",
        "wf_fixed_base_to_kaiku_stg",
        "wf_initKaikuMeta",
        "wf_InitPartyIdLkp",
        "wf_INVOICE_to_kaiku_stg",
        "wf_InvoiceData_ToKaikuSTG",
        "wf_JPD_to_kaiku_stg",
        "wf_kaiku_datamart_report",
        "wf_kaiku_telco_cloud",
        "wf_mankeli_hierarchy_to_kaiku_stg",
        "wf_mankeli_to_kaiku_stg",
        "wf_nirap_tables_to_kaiku_stg",
        "wf_odw_busToKaikuSTG",
        "wf_ODW_to_kaiku_stg",
        "wf_Prepaid",
        "wf_Rambo",
        "wf_Roaming",
        "wf_sales_to_kaiku_stg",
        "wf_salesforce_kaiku_stg",
        "wf_SMS_IN_workaround",
        "wf_TERRA_to_kaiku_stg",
        "wf_UsageIC",
        "wf_VAS"
    ]

    try:
        col_map = CONFIG["COLUMN_MAPPINGS"]
        print(f'col_map : {col_map}')
        print(f"SCHEDULE_DETAILS_PATH = {CONFIG['SCHEDULE_DETAILS_PATH']}")
        df_schedule = pd.read_csv(CONFIG["SCHEDULE_DETAILS_PATH"], encoding="utf-8-sig")
        df_schedule.columns = df_schedule.columns.str.strip()
        df_task_config = pd.read_csv(CONFIG["TASK_CONFIG_PATH"], encoding="utf-8-sig")
        df_task_config.columns = df_task_config.columns.str.strip()
        df_source_target = pd.read_csv( CONFIG["SOURCE_TARGET_PATH"], encoding="utf-8-sig")
        print(f"df_source_target : {df_source_target}")
        df_source_target.columns = df_source_target.columns.str.strip()
        print(f'df_source_target.columns : {df_source_target.columns}')
        df_dag_config = pd.read_csv(CONFIG["DAG_CONFIG_PATH"], encoding="utf-8-sig")
        df_dag_config.columns = df_dag_config.columns.str.strip()
        df_overview = pd.read_csv(CONFIG["PROCESS_OVERVIEW_PATH"], encoding="utf-8-sig")
        df_overview.columns = df_overview.columns.str.strip()
        dataframes = {
            "schedule": df_schedule.set_index(col_map["WORKFLOW_NAME"]),
            "source_target": df_source_target,
            "dag_config": df_dag_config,
            "task_config": df_task_config.set_index(col_map["WORKFLOW_NAME"]),
            "overview": df_overview,
        }
        workflow_col = col_map["DAG_NAME_LOOKUP"]
        print(f'workflow_col : {workflow_col}')
        dataframes["source_target"][workflow_col] = dataframes["source_target"][
            workflow_col
        ].ffill()
        dataframes["dag_config"][workflow_col] = dataframes["dag_config"][workflow_col].ffill()
        for wf_name in workflows_to_generate:
            print(f"wf_name is : {wf_name}")
            print(f"\n--- Generating runbook for {wf_name} ---")
            wf_details = get_all_workflow_details(wf_name, dataframes, CONFIG)
            print(f"wf_details")
            if wf_details:
                create_runbook_doc(wf_name, wf_details, CONFIG)

    except FileNotFoundError as e:
        print(f"❌ ERROR: The file was not found: {e.filename}")
        print("👉 Please ensure the path in the CONFIG section is correct.")
    except KeyError as e:
        print(f"❌ ERROR: A required column was not found: {e}.")
        print(
            f"👉 Please check the column name for {e} in the 'COLUMN_MAPPINGS' section of the script and ensure it matches your CSV file."
        )
    except Exception as e:
        print(f"❌ An unexpected error occurred: {e}")


